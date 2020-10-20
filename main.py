import argparse
import math
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches  # 英寸

path = ''
linesum = 3
exportpath = ''

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def start_genarate():
    piclists = os.listdir(path)
    print(piclists)
    # exit()
    document = Document()
    rownums = int(math.ceil(len(piclists) / linesum))

    print('设表格行', rownums)
    if len(piclists) != 0:
        # 插入表格   表格是从1，0开始，第一行就是1，列是从0开始，NND
        table = document.add_table(rows=rownums, cols=linesum, style='Table Grid')
        for rownum in range(rownums):
            for cellnum in range(linesum):

                cell = table.cell(rownum, cellnum)
                p = cell.paragraphs[0]

                run = p.add_run()
                picnum = (rownum) * linesum + cellnum
                if picnum < len(piclists):
                    run.add_picture(os.path.join(path, piclists[picnum]), width=Inches(1.25))  # ,width=Cm(jpgwid))
                    # print('写入',str(rownum),"0",os.path.join(path1,jpglists[int((rownum)*2)]))
                    set_cell_border(cell,
                                    top={"sz": 12, "val": "single", "color": "FFFFFF", "space": "0"},
                                    bottom={"sz": 12, "color": "FFFFFF", "val": "single"},
                                    left={"sz": 12, "val": "dashed", "color": "FFFFFF", "shadow": "true"},
                                    right={"sz": 10, "color": "FFFFFF", "val": "dashed"},
                                    insideH={"color": "FFFFFF"}, )
                elif (picnum % linesum) < linesum:
                    set_cell_border(cell,
                                    top={"sz": 12, "val": "single", "color": "FFFFFF", "space": "0"},
                                    bottom={"sz": 12, "color": "FFFFFF", "val": "single"},
                                    left={"sz": 12, "val": "dashed", "color": "FFFFFF", "shadow": "true"},
                                    right={"sz": 10, "color": "FFFFFF", "val": "dashed"},
                                    insideH={"color": "FFFFFF"}, )
        document.save(exportpath)  # 保存文档


if __name__ == "__main__":
    parser = argparse.ArgumentParser()

    parser.add_argument('-p', '--path', type=str, default=None, required=True, help='要导入的图片路径')
    parser.add_argument('-l', '--linesum', type=int, default=3, help='每一行的图片个数, 默认为：3')
    parser.add_argument('-e', '--exportpath', type=str, default="pic.docx", help='导出文件所在路径，默认为当前路径下：pic.docx')

    args = parser.parse_args()
    path = args.path
    linesum = args.linesum
    exportpath = args.exportpath

    start_genarate()

